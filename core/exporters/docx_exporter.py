"""Markdown -> DOCX exporter using python-docx.

This exporter is the lightweight, pip-native DOCX renderer. It does not call
Pandoc, Word, LibreOffice, or an external binary. Its job is final rendering:
preserve validated Markdown structure and apply a conservative public-document
style profile without inventing or rewriting content.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

from .export_base import BaseExporter, ExportResult
from .markdown_blocks import Heading, ListBlock, Paragraph, Run, Table, parse_markdown
from .style_profile import DEFAULT_PUBLIC_DOCUMENT_STYLE_PROFILE, DocumentStyleProfile


class DocxExporter(BaseExporter):
    """Markdown file -> DOCX document."""

    supported_ext = (".docx",)

    def __init__(self, style_profile: DocumentStyleProfile | None = None) -> None:
        self.style_profile = style_profile or DEFAULT_PUBLIC_DOCUMENT_STYLE_PROFILE

    def export(self, markdown_path: Path, output_path: Path) -> ExportResult:
        markdown_path = Path(markdown_path)
        output_path = Path(output_path)

        if not self.can_export(output_path):
            return ExportResult(
                source=markdown_path,
                output=output_path,
                ok=False,
                error=(
                    f"Unsupported output extension: {output_path.suffix} "
                    f"(supported: {sorted(self.supported_ext)})"
                ),
                meta={"exporter": self.name, "requires_optional_tool": False},
                error_code="export_unsupported_extension",
            )
        if not markdown_path.exists():
            return ExportResult(
                source=markdown_path,
                output=output_path,
                ok=False,
                error=f"Markdown source does not exist: {markdown_path}",
                meta={"exporter": self.name, "requires_optional_tool": False},
                error_code="export_source_missing",
            )

        try:
            blocks = parse_markdown(markdown_path.read_text(encoding="utf-8"))
            document = Document()
            self._apply_style_profile(document)

            table_stats = self._table_stats(blocks)
            self._apply_table_layout_strategy(document, table_stats)

            for block in blocks:
                self._write_block(document, block)

            output_path.parent.mkdir(parents=True, exist_ok=True)
            document.save(str(output_path))
            return ExportResult(
                source=markdown_path,
                output=output_path,
                ok=True,
                meta={
                    "exporter": self.name,
                    "blocks": len(blocks),
                    "requires_optional_tool": False,
                    "format": ".docx",
                    "style_profile": self.style_profile.profile_id,
                    "docx_quality_level": "partially_stabilized",
                    **table_stats,
                    "note": (
                        "pip-native DOCX export; public-document style, "
                        "visible text, headings, lists, and tables are tested; "
                        "layout-perfect output is not claimed"
                    ),
                },
            )
        except Exception as e:  # noqa: BLE001 - exporters return structured errors.
            return ExportResult(
                source=markdown_path,
                output=output_path,
                ok=False,
                error=repr(e),
                meta={"exporter": self.name, "requires_optional_tool": False},
                error_code="export_failed",
            )

    # --- Block rendering -------------------------------------------------

    def _write_block(self, document, block) -> None:  # noqa: ANN001
        if isinstance(block, Heading):
            self._write_heading(document, block)
        elif isinstance(block, Paragraph):
            self._write_paragraph(document, block)
        elif isinstance(block, ListBlock):
            self._write_list(document, block)
        elif isinstance(block, Table):
            self._write_table(document, block)

    def _write_heading(self, document, block: Heading) -> None:  # noqa: ANN001
        level = min(max(block.level, 1), 9)
        paragraph = document.add_paragraph(style=f"Heading {level}")
        self._add_runs(paragraph, block.runs, bold_override=True)
        self._format_heading_paragraph(paragraph, level)

    def _write_paragraph(self, document, block: Paragraph) -> None:  # noqa: ANN001
        paragraph = document.add_paragraph()
        self._add_runs(paragraph, block.runs)
        self._format_body_paragraph(paragraph)

    def _write_list(self, document, block: ListBlock) -> None:  # noqa: ANN001
        style = "List Number" if block.ordered else "List Bullet"
        for item in block.items:
            paragraph = document.add_paragraph(style=style)
            self._add_runs(paragraph, item)
            self._format_list_paragraph(paragraph)

    def _write_table(self, document, block: Table) -> None:  # noqa: ANN001
        body_cols = max((len(row) for row in block.rows), default=0)
        ncols = max(len(block.header), body_cols)
        if ncols == 0:
            return

        rows = block.rows
        nrows = (1 if block.header else 0) + len(rows)
        if nrows == 0:
            return

        table = document.add_table(rows=nrows, cols=ncols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_fixed_table_layout(table)
        self._set_table_width(table, width_pct=100)

        row_index = 0
        if block.header:
            self._fill_row(table.rows[0].cells, block.header, ncols, ncols, is_header=True)
            self._set_cant_split_row(table.rows[0])
            row_index = 1

        for row in rows:
            self._fill_row(table.rows[row_index].cells, row, ncols, ncols, is_header=False)
            self._set_cant_split_row(table.rows[row_index])
            row_index += 1

        self._apply_column_widths(document, table, ncols)
        self._format_table_cells(table, ncols)
        document.add_paragraph()

    def _fill_row(
        self,
        cells,
        row_runs: list[list[Run]],
        ncols: int,
        table_cols: int,
        *,
        is_header: bool,
    ) -> None:  # noqa: ANN001
        font_size = self._table_font_size(table_cols)
        for col_index in range(ncols):
            runs = row_runs[col_index] if col_index < len(row_runs) else []
            paragraph = cells[col_index].paragraphs[0]
            self._add_runs(
                paragraph,
                runs,
                font_size_pt=font_size,
                bold_override=is_header,
            )
            if is_header:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._shade_cell(cells[col_index], self.style_profile.table_header_fill)

    def _add_runs(
        self,
        paragraph,
        runs: list[Run],
        font_size_pt: float | None = None,
        *,
        bold_override: bool = False,
    ) -> None:  # noqa: ANN001
        for run in runs:
            rendered = paragraph.add_run(run.text)
            rendered.bold = True if bold_override else run.bold
            rendered.italic = run.italic
            rendered.font.name = self.style_profile.font_family
            if font_size_pt is not None:
                rendered.font.size = Pt(font_size_pt)
            self._set_run_eastasia_font(rendered, self.style_profile.font_family)

    # --- Style application ----------------------------------------------

    def _apply_style_profile(self, document) -> None:  # noqa: ANN001
        """Apply the project-local DOCX style profile."""
        profile = self.style_profile

        for section in document.sections:
            section.page_width = Mm(profile.page_width_mm)
            section.page_height = Mm(profile.page_height_mm)
            section.top_margin = Mm(profile.page_margin_top_mm)
            section.bottom_margin = Mm(profile.page_margin_bottom_mm)
            section.left_margin = Mm(profile.page_margin_left_mm)
            section.right_margin = Mm(profile.page_margin_right_mm)

        normal = document.styles["Normal"]
        normal.font.name = profile.font_family
        normal.font.size = Pt(profile.font_size_pt)
        self._set_style_eastasia_font(normal, profile.font_family)
        normal_format = normal.paragraph_format
        normal_format.line_spacing = profile.line_spacing
        normal_format.space_after = Pt(profile.paragraph_space_after_pt)

        self._configure_heading_style(document, "Heading 1", profile.heading_font_size_pt)
        heading1_alignment = self._alignment(profile.heading_alignment)
        if heading1_alignment is not None:
            document.styles["Heading 1"].paragraph_format.alignment = heading1_alignment
        self._configure_heading_style(document, "Heading 2", profile.heading2_font_size_pt)
        document.styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._configure_heading_style(document, "Heading 3", profile.heading3_font_size_pt)
        document.styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for level in range(4, 10):
            self._configure_heading_style(document, f"Heading {level}", profile.font_size_pt)
            document.styles[f"Heading {level}"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _configure_heading_style(self, document, style_name: str, size_pt: float) -> None:  # noqa: ANN001
        style = document.styles[style_name]
        style.font.name = self.style_profile.font_family
        style.font.size = Pt(size_pt)
        style.font.bold = True
        self._set_style_eastasia_font(style, self.style_profile.font_family)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    def _format_heading_paragraph(self, paragraph, level: int) -> None:  # noqa: ANN001
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.line_spacing = 1.15
        if level == 1:
            paragraph.alignment = self._alignment(self.style_profile.heading_alignment)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(14)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(6)

    def _format_body_paragraph(self, paragraph) -> None:  # noqa: ANN001
        fmt = paragraph.paragraph_format
        fmt.line_spacing = self.style_profile.line_spacing
        fmt.space_after = Pt(self.style_profile.paragraph_space_after_pt)
        if self.style_profile.body_first_line_indent_mm:
            fmt.first_line_indent = Mm(self.style_profile.body_first_line_indent_mm)
        if paragraph.text.strip():
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _format_list_paragraph(self, paragraph) -> None:  # noqa: ANN001
        fmt = paragraph.paragraph_format
        fmt.left_indent = Mm(8)
        fmt.first_line_indent = Mm(-4)
        fmt.line_spacing = self.style_profile.line_spacing
        fmt.space_after = Pt(3)

    # --- Table quality helpers ------------------------------------------

    def _table_stats(self, blocks: list) -> dict:
        table_blocks = [block for block in blocks if isinstance(block, Table)]
        column_counts = [
            max(len(table.header), max((len(row) for row in table.rows), default=0))
            for table in table_blocks
        ]
        max_cols = max(column_counts, default=0)
        wide = max_cols >= 8
        strategy = "landscape_compact_table" if wide else "fixed_width_table_grid"
        warnings = []
        if wide:
            warnings.append(
                "wide_table_detected: landscape orientation, compact margins, "
                "fixed table layout, and smaller table font were applied"
            )
        return {
            "table_count": len(table_blocks),
            "max_table_column_count": max_cols,
            "wide_table_detected": wide,
            "wide_table_strategy": strategy,
            "warnings": warnings,
        }

    def _apply_table_layout_strategy(self, document, table_stats: dict) -> None:  # noqa: ANN001
        if not table_stats["wide_table_detected"]:
            return
        for section in document.sections:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Mm(self.style_profile.page_height_mm)
            section.page_height = Mm(self.style_profile.page_width_mm)
            section.left_margin = Mm(12)
            section.right_margin = Mm(12)
            section.top_margin = Mm(18)
            section.bottom_margin = Mm(15)

    def _table_font_size(self, ncols: int) -> float:
        if ncols >= 10:
            return 7.5
        if ncols >= 8:
            return 8.0
        if ncols >= 6:
            return 9.0
        return self.style_profile.table_font_size_pt

    def _apply_column_widths(self, document, table, ncols: int) -> None:  # noqa: ANN001
        if ncols <= 0:
            return
        usable_width_mm = self._usable_width_mm(document)
        col_width_mm = max(usable_width_mm / ncols, 8)
        for row in table.rows:
            for cell in row.cells:
                cell.width = Mm(col_width_mm)
                self._set_cell_width(cell, col_width_mm)

    def _usable_width_mm(self, document) -> float:  # noqa: ANN001
        section = document.sections[0]
        return max(section.page_width.mm - section.left_margin.mm - section.right_margin.mm, 20)

    def _format_table_cells(self, table, ncols: int) -> None:  # noqa: ANN001
        font_size = self._table_font_size(ncols)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                margin = self.style_profile.table_cell_margin_twips
                self._set_cell_margins(cell, top=margin, bottom=margin, start=margin, end=margin)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    if row_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = self.style_profile.font_family
                        run.font.size = Pt(font_size)
                        if row_index == 0:
                            run.bold = True
                        self._set_run_eastasia_font(run, self.style_profile.font_family)

    @staticmethod
    def _set_fixed_table_layout(table) -> None:  # noqa: ANN001
        tbl_pr = table._tbl.tblPr  # noqa: SLF001 - python-docx table XML access.
        tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")

    @staticmethod
    def _set_table_width(table, width_pct: int) -> None:  # noqa: ANN001
        tbl_pr = table._tbl.tblPr  # noqa: SLF001
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
            tbl_pr.append(tbl_w)
        tbl_w.set(qn("w:type"), "pct")
        tbl_w.set(qn("w:w"), str(width_pct * 50))

    @staticmethod
    def _set_cell_width(cell, width_mm: float) -> None:  # noqa: ANN001
        tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:type"), "dxa")
        tc_w.set(qn("w:w"), str(int(width_mm / 25.4 * 1440)))

    @staticmethod
    def _set_cant_split_row(row) -> None:  # noqa: ANN001
        tr_pr = row._tr.get_or_add_trPr()  # noqa: SLF001
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    @staticmethod
    def _shade_cell(cell, fill: str) -> None:  # noqa: ANN001
        tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)

    @staticmethod
    def _set_cell_margins(cell, top: int, bottom: int, start: int, end: int) -> None:  # noqa: ANN001
        tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001
        tc_mar = tc_pr.first_child_found_in("w:tcMar")
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)

        for margin_name, value in {
            "top": top,
            "bottom": bottom,
            "start": start,
            "end": end,
        }.items():
            node = tc_mar.find(qn(f"w:{margin_name}"))
            if node is None:
                node = OxmlElement(f"w:{margin_name}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    # --- Font and alignment helpers -------------------------------------

    @staticmethod
    def _set_style_eastasia_font(style, font_name: str) -> None:  # noqa: ANN001
        rpr = style.element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _set_run_eastasia_font(run, font_name: str) -> None:  # noqa: ANN001
        rpr = run._element.get_or_add_rPr()  # noqa: SLF001
        rpr.get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    @staticmethod
    def _alignment(name: str):
        return {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(name.lower())
