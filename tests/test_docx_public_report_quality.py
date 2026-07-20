"""DOCX public-report rendering quality regression checks."""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from core.exporters.docx_exporter import DocxExporter
from core.exporters.style_profile import DEFAULT_PUBLIC_DOCUMENT_STYLE_PROFILE


REPORT_MD = """\
# 2026년 공공기관 안전관리 추진계획

수신: 확인 필요
담당: 안전관리팀

## 1. 추진 배경

공공기관 안전관리 수준을 점검하고 개선 과제를 체계적으로 관리하기 위하여 추진계획을 수립한다.

## 2. 주요 추진 과제

1. 안전관리 체계 점검
2. 취약시설 개선 계획 수립
3. 이행 현황 점검 및 결과 보고

| 구분 | 주요 내용 | 담당 | 기한 |
| --- | --- | --- | --- |
| 체계 점검 | 안전관리 기준과 현장 이행 상태를 비교 점검 | 안전관리팀 | 2026. 7. 15. |
| 개선 계획 | 취약 항목별 개선 일정과 담당자를 명확히 지정 | 시설관리팀 | 확인 필요 |

붙임  세부 추진계획 1부.  끝.
"""


def _all_text(document: Document) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _table_has_fixed_layout(table) -> bool:  # noqa: ANN001
    layout = table._tbl.tblPr.first_child_found_in("w:tblLayout")  # noqa: SLF001
    return layout is not None and layout.get(qn("w:type")) == "fixed"


def _cell_fill(cell) -> str | None:  # noqa: ANN001
    shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))  # noqa: SLF001
    if shading is None:
        return None
    return shading.get(qn("w:fill"))


def test_docx_public_report_style_and_table_rendering() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        md_path = tmp_path / "public_plan.md"
        docx_path = tmp_path / "public_plan.docx"
        md_path.write_text(REPORT_MD, encoding="utf-8")

        result = DocxExporter().export(md_path, docx_path)

        assert result.ok, result.error
        assert result.meta["docx_quality_level"] == "partially_stabilized"
        assert result.meta["requires_optional_tool"] is False
        assert docx_path.exists() and docx_path.stat().st_size > 0

        document = Document(str(docx_path))
        text = _all_text(document)
        for token in (
            "2026년 공공기관 안전관리 추진계획",
            "공공기관 안전관리 수준",
            "안전관리 체계 점검",
            "2026. 7. 15.",
            "붙임",
            "끝.",
        ):
            assert token in text

        heading1 = document.paragraphs[0]
        assert heading1.style.name == "Heading 1"
        assert heading1.alignment == WD_ALIGN_PARAGRAPH.CENTER
        assert all(run.bold for run in heading1.runs if run.text)

        normal = document.styles["Normal"]
        assert normal.font.name == DEFAULT_PUBLIC_DOCUMENT_STYLE_PROFILE.font_family
        assert abs(normal.font.size.pt - DEFAULT_PUBLIC_DOCUMENT_STYLE_PROFILE.font_size_pt) < 0.1

        assert len(document.tables) == 1
        table = document.tables[0]
        assert _table_has_fixed_layout(table)
        assert table.cell(0, 0).text == "구분"
        assert table.cell(0, 1).text == "주요 내용"
        assert _cell_fill(table.cell(0, 0)) == DEFAULT_PUBLIC_DOCUMENT_STYLE_PROFILE.table_header_fill
        assert all(run.bold for run in table.cell(0, 0).paragraphs[0].runs if run.text)


if __name__ == "__main__":
    test_docx_public_report_style_and_table_rendering()
    print("PASS: DOCX public report quality")
