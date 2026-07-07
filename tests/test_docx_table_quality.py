"""DOCX form/table quality regression checks."""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document
from docx.enum.section import WD_ORIENT

from core.exporters.docx_exporter import DocxExporter


FIXTURES = [
    Path(__file__).resolve().parent / "fixtures" / "export" / "wide_activity_report.md",
    Path(__file__).resolve().parent / "fixtures" / "export" / "business_plan_form.md",
]


def _all_text(document: Document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_docx_preserves_form_tables_and_long_cells() -> None:
    for fixture in FIXTURES:
        with tempfile.TemporaryDirectory() as tmp:
            out = Path(tmp) / f"{fixture.stem}.docx"
            result = DocxExporter().export(fixture, out)

            assert result.ok, result.error
            assert out.exists() and out.stat().st_size > 0
            assert result.meta["table_count"] >= 3
            assert result.meta["max_table_column_count"] >= 8
            assert result.meta["wide_table_detected"] is True
            assert result.meta["wide_table_strategy"] == "landscape_compact_table"

            document = Document(str(out))
            text = _all_text(document)

            assert len(document.tables) >= 3
            assert any(len(table.columns) >= 8 for table in document.tables)
            assert "확인 필요" in text
            assert "실제 개인정보를 넣지 않는다" in text or "충분히 긴 한국어 문장" in text
            assert document.sections[0].orientation == WD_ORIENT.LANDSCAPE
            assert document.sections[0].page_width > document.sections[0].page_height


if __name__ == "__main__":
    test_docx_preserves_form_tables_and_long_cells()
    print("PASS: DOCX table quality")
