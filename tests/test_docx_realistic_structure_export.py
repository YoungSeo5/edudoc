"""Loop 8.99: DOCX structure regression on a realistic wide-table document.

Goes beyond "file exists": checks that tables, wide-table columns, long-cell text,
headings, and the style profile survive export. Catches the shallow-smoke-test gap.
Does NOT assert exact Microsoft Word rendering or layout-perfect preservation.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document

from core.exporters.docx_exporter import DocxExporter
from core.exporters.style_profile import DEFAULT_PUBLIC_DOCUMENT_STYLE_PROFILE

FIXTURE = (
    Path(__file__).resolve().parent
    / "fixtures" / "export" / "wide_table_activity_report.md"
)


def _all_text(document) -> str:  # noqa: ANN001
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_docx_realistic_structure_export() -> None:
    assert FIXTURE.exists(), f"fixture missing: {FIXTURE}"

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "report.docx"
        result = DocxExporter().export(FIXTURE, out)
        assert result.ok, result.error
        assert out.exists() and out.stat().st_size > 0, "DOCX missing/empty"

        document = Document(str(out))
        text = _all_text(document)

        # content-preserving: key text + a long table-cell sentence survive
        for token in (
            "중간 활동보고서",
            "예시모둠",
            "활동 성찰",
            "자격증 대비 문제 풀이를 병행",  # long-cell text
        ):
            assert token in text, f"missing content: {token}"

        # structure-preserving: multiple tables incl. a wide (>=8 col) one
        assert len(document.tables) >= 3, f"expected >=3 tables, got {len(document.tables)}"
        assert any(len(t.columns) >= 8 for t in document.tables), "no wide (>=8 col) table"

        # headings recognizable
        headings = [p for p in document.paragraphs if p.style.name.startswith("Heading")]
        assert any(
            "활동보고서" in p.text or "일반 현황" in p.text for p in headings
        ), "section headings not recognizable"

        # style profile still applied (pip-native, no external office tool)
        normal = document.styles["Normal"]
        assert normal.font.size is not None
        assert abs(normal.font.size.pt - DEFAULT_PUBLIC_DOCUMENT_STYLE_PROFILE.font_size_pt) < 0.1


if __name__ == "__main__":
    test_docx_realistic_structure_export()
    print("PASS: DOCX realistic structure export")
