"""
DocxExporter 검증 (docs/exporter-docx-success.md 의 필수 항목).

깨끗한 공문형 Markdown 초안 -> DOCX 왕복이 pip-native(python-docx)로
외부 바이너리 없이 도는지 확인한다. 통과하면 이 슬라이스가 다른 exporter의
템플릿이 된다.
"""
from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document

from core.exporters.docx_exporter import DocxExporter

SAMPLE_MD = """\
# 공고문 제목

관련 근거에 따라 아래와 같이 **안내**합니다.

1. 첫째 항목
2. 둘째 항목

| 구분 | 내용 |
| --- | --- |
| 성명 | 홍길동 |
| 소속 | 총무과 |

붙임  안내문 1부.  끝.
"""


def main() -> int:
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        md_path = tmp / "draft.md"
        out_path = tmp / "draft.docx"
        md_path.write_text(SAMPLE_MD, encoding="utf-8")

        result = DocxExporter().export(md_path, out_path)
        assert result.ok, f"export 실패: {result.error}"

        # 1. 파일 생성 + 유효한 docx(zip 매직 PK)
        assert out_path.exists() and out_path.stat().st_size > 0, "출력 파일 없음/빈 파일"
        assert out_path.read_bytes()[:2] == b"PK", "유효한 docx(zip)가 아님"

        doc = Document(str(out_path))

        # 2. 제목 보존 (Heading 스타일 + 텍스트)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("공고문 제목" in p.text for p in headings), "제목이 Heading 으로 안 남음"

        # 3. 목록 보존 (List Number 항목 2개)
        numbered = [p for p in doc.paragraphs if p.style.name == "List Number"]
        assert len(numbered) == 2, f"순서 목록 항목 2개 기대, 실제 {len(numbered)}"
        assert numbered[0].text == "첫째 항목"

        # 4. 표 보존 (3행 2열 + 셀 텍스트)
        assert len(doc.tables) == 1, "표 1개 기대"
        table = doc.tables[0]
        assert len(table.rows) == 3 and len(table.columns) == 2, "표 크기 3x2 불일치"
        assert table.cell(0, 0).text == "구분"
        assert table.cell(1, 1).text == "홍길동"

        # 5. 서식 보존 (bold run 존재)
        bold_texts = [
            r.text for p in doc.paragraphs for r in p.runs if r.bold
        ]
        assert any("안내" in t for t in bold_texts), "굵게 서식이 안 남음"

        # 6. 한글 무결 (왕복 텍스트 일치)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        for token in ("공고문 제목", "안내", "붙임", "끝."):
            assert token in all_text, f"한글 텍스트 누락: {token}"

        print("PASS: DocxExporter 왕복(md->docx) 정상")
        print("  제목:", [p.text for p in headings])
        print("  목록:", [p.text for p in numbered])
        print("  표:", f"{len(table.rows)}x{len(table.columns)}",
              [table.cell(0, 0).text, table.cell(0, 1).text])
        return 0


if __name__ == "__main__":
    raise SystemExit(main())
